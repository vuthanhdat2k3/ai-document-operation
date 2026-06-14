from __future__ import annotations

import logging
from pathlib import Path

from docx import Document as DocxDocument
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph

from app.processing.parsers.base import (
    BaseParser,
    ImageData,
    PageResult,
    ParseResult,
    TableData,
)

logger = logging.getLogger(__name__)

_HEADING_PATTERN = {"Heading 1", "Heading 2", "Heading 3", "Heading 4",
                    "Heading 5", "Heading 6", "Title", "Subtitle"}


class DOCXParser(BaseParser):
    """Parse DOCX files using python-docx.

    Capabilities:
    - Paragraph extraction with heading-level metadata
    - Table extraction (rows and cells)
    - Core-properties metadata extraction
    """

    def parse(self, file_path: Path) -> ParseResult:
        if not file_path.exists():
            raise FileNotFoundError(f"DOCX not found: {file_path}")

        try:
            doc = DocxDocument(str(file_path))
        except Exception as exc:
            raise ValueError(f"Cannot open DOCX: {exc}") from exc

        text_parts: list[str] = []
        tables: list[TableData] = []
        images: list[ImageData] = []

        try:
            images = self._extract_images(doc)
        except Exception:
            logger.warning("Image extraction failed", exc_info=True)

        try:
            tables = self._extract_tables(doc)
        except Exception:
            logger.warning("Table extraction failed", exc_info=True)

        try:
            text_parts = self._extract_paragraphs(doc)
        except Exception:
            logger.warning("Paragraph extraction failed", exc_info=True)

        full_text = "\n".join(text_parts)
        table_text = "\n".join(
            " | ".join(row) for t in tables for row in [t.headers] + t.rows
        )
        combined_text = f"{full_text}\n{table_text}".strip()

        metadata = self._extract_metadata(doc)

        char_count = len(combined_text)
        confidence = min(1.0, char_count / 500) if char_count < 500 else 1.0

        page_result = PageResult(
            page_number=1,
            text=combined_text,
            tables=tables,
            images=images,
            confidence=round(confidence, 4),
        )

        return ParseResult(
            pages=[page_result],
            metadata=metadata,
            quality_score=round(confidence, 4),
        )

    def _extract_paragraphs(self, doc: DocxDocument) -> list[str]:
        parts: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else ""
            if style_name in _HEADING_PATTERN:
                level = self._heading_level(style_name)
                prefix = "#" * level + " " if level else ""
                parts.append(f"{prefix}{text}")
            else:
                parts.append(text)
        return parts

    @staticmethod
    def _heading_level(style_name: str) -> int:
        if style_name.startswith("Heading "):
            try:
                return int(style_name.split()[-1])
            except (ValueError, IndexError):
                return 0
        if style_name == "Title":
            return 1
        if style_name == "Subtitle":
            return 2
        return 0

    def _extract_tables(self, doc: DocxDocument) -> list[TableData]:
        result: list[TableData] = []
        for idx, table in enumerate(doc.tables):
            try:
                td = self._parse_table(table, idx + 1)
                result.append(td)
            except Exception:
                logger.warning("Failed to parse table %d", idx + 1, exc_info=True)
        return result

    @staticmethod
    def _parse_table(table: DocxTable, table_idx: int) -> TableData:
        rows_data: list[list[str]] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(cells)

        if not rows_data:
            return TableData(headers=[], rows=[], page=table_idx)

        headers = rows_data[0]
        data_rows = rows_data[1:] if len(rows_data) > 1 else []
        return TableData(headers=headers, rows=data_rows, page=table_idx)

    @staticmethod
    def _extract_images(doc: DocxDocument) -> list[ImageData]:
        images: list[ImageData] = []
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    images.append(
                        ImageData(
                            description=f"Embedded image: {rel.target_ref}",
                            storage_path=None,
                        )
                    )
        except Exception:
            logger.debug("Could not enumerate DOCX image relationships", exc_info=True)
        return images

    @staticmethod
    def _extract_metadata(doc: DocxDocument) -> dict:
        try:
            props = doc.core_properties
        except Exception:
            return {}

        return {
            "title": props.title or "",
            "author": props.author or "",
            "subject": props.subject or "",
            "keywords": props.keywords or "",
            "created": str(props.created) if props.created else "",
            "modified": str(props.modified) if props.modified else "",
            "last_modified_by": props.last_modified_by or "",
            "revision": props.revision or 0,
        }
